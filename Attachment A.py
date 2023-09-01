# -*- coding: utf-8 -*-
"""
Created on Tue May 11, 2023

@author: rohiniiyengar
Purpose: Create an Attachment A Word document that contains Quickbase data related to the specified application number.
"""
import math
import docx
import pandas as pd
import json
import requests
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxcompose.composer import Composer
from docx.oxml.shared import OxmlElement, qn
import base64
from pypdf import PdfMerger

downloadFolder = r'C:\Users\58741\Downloads'
destinationFolder = r'C:\Users\58741\OneDrive - ICF\Documents\Python Scripts\Attachment A'
finalDestination = r'C:\Users\58741\OneDrive - ICF\Documents\Python Scripts\Attachment A\Apps for Contract\08.15.2023'
usertoken = 'b7n73q_g6xg_0_ttt8jjdvzpzzudurmt9vcryr5yz'

os.chdir(destinationFolder)

def readQBData(tableid, reportid, filename, verbose=True):
    headers = {
        'QB-Realm-Hostname': 'dmsrecovery.quickbase.com',
        'Authorization': 'QB-USER-TOKEN ' + 'b7n73q_g6xg_0_ttt8jjdvzpzzudurmt9vcryr5yz'
    }
    params = {'tableId': tableid}
    r = requests.post('https://api.quickbase.com/v1/reports/' + reportid + '/run', params=params, headers=headers)

    # get data, unnet
    qb_data = pd.json_normalize(r.json()['data'])

    # get the field names and IDs as a dict
    qb_fields = pd.json_normalize(r.json()['fields'])
    qb_fields_dict = dict(zip(qb_fields.id.astype(str), qb_fields.label))

    # look up fids in dict to get correctly ordered list of field names
    new_cols = []
    qb_data.columns = qb_data.columns.str.strip('.value')
    for colName in qb_fields_dict.items():
        new_cols.append(colName[0])
    qb_data = qb_data[new_cols]

    # map column ids to their name
    qb_data.rename(columns=qb_fields_dict, inplace=True)

    qb_data.to_csv(filename + '.csv')

    return qb_data

def readIntake(Intake):
    os.chdir(destinationFolder)
    data = readQBData('bsa4kcifd', '76', "Attachment A: Intake")
    data = data[data['Record ID#'] == Intake]
    data.to_csv(destinationFolder + "\Attachment A Intake.csv", sep='|', index=False)
    return data

def readActivity(Intake):
    os.chdir(destinationFolder)
    data = readQBData('bsbtxz7z2', '10', "Attachment A: Activity")
    data = data[data['Related Intake'] == Intake]
    data.to_csv(destinationFolder + "\Attachment A Activity.csv", sep='|', index=False)
    return data

def readProject(Intake, Activity):
    os.chdir(destinationFolder)
    data = readQBData('bsbtx4cu6', '26', "Attachment A: Project")
    data = data[data['Related Intake'] == Intake]
    data = data[data['Related Activity'] == Activity]
    data.to_csv(destinationFolder + "\Attachment A Project.csv", sep='|', index=False)
    return data

def readProjectSite(Intake, Activity, Project):
    os.chdir(destinationFolder)
    data = readQBData('bsbui5m7t', '13', "Attachment A: Project Sites")
    data = data[data['Project - Related Intake'] == Intake]
    data = data[data['Project - Related Activity'] == Activity]
    data = data[data['Related Project'] == Project]
    data.to_csv(destinationFolder + "\Attachment A Project Sites.csv", sep='|', index=False)
    return data

def set_line_spacing(document):
    for paragraph in document.paragraphs:
        para_format = paragraph.paragraph_format
        para_format.space_after = 0

def preventDocumentBreak(document):
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]                     # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)                   # Append in the new tag

def createDoc(Intake):
    intake = readIntake(Intake)

    # create doc for a specific intake, loop through each activity, loop through each project (where related project = record id)

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    font_styles = document.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(14)
    font_object.name = 'Times New Roman'
    font_object.bold = True

    heading1 = document.add_paragraph()
    heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading1.add_run(intake['Applicant: Applicant'].values[0], style='CommentsStyle')

    heading2 = document.add_paragraph()
    heading2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading2.add_run("##-###-000-X###", style='CommentsStyle')

    heading3 = document.add_paragraph()
    heading3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading3.add_run("PERFORMANCE STATEMENT", style='CommentsStyle')
    document.add_paragraph()

    def DRGRActivity():
        string = ""
        for i in range(len(intake['DRGR Activity'].values[0])):
            string = intake['DRGR Activity'].values[0][0]
            if i > 0:
                string += " and " + intake['DRGR Activity'].values[0][i]
        return string

    word = ""
    if "City" in intake['Applicant: Applicant'].values[0]:
        word = "the "

    para1 = document.add_paragraph(
        "The GLO awards " + word + intake['Applicant: Applicant'].values[0] + " (Subrecipient) this Contract under HUD’s Community Development Block Grant Mitigation "
        "(“CDBG-MIT”) program to provide financial assistance with funds appropriated to facilitate Activities related to disaster relief, long-term recovery, restoration "
        "of infrastructure and housing, economic revitalization, mitigation, and affirmatively furthering fair housing, in accordance with Executive Order 12892, in the most "
        "impacted and distressed areas resulting from a major declared disaster that occurred in 2015, 2016, 2017, or 2018.")
    document.add_paragraph()

    para2 = document.add_paragraph(
        "In strict conformance with the terms and conditions of the CDBG-MIT – 2015/2016/Hurricane Harvey " + intake['State/HUD MID'].values[0] + " MID Regional Mitigation Program and this Contract, "
        "Subrecipient shall perform, or cause to be performed, the Infrastructure Activities identified below to increase its resilience to disasters and reduce or eliminate "
        "long-term risk of disaster-related loss of life, injury, damage to and loss of property, and suffering and hardship by lessening the impact of future disasters.")
    document.add_paragraph()

    para3 = document.add_paragraph(
        "Subrecipient shall perform the Activities identified herein for the service area specified in its approved Texas Community "
        "Development Block Grant Mitigation Grant Application to provide a long-lasting investment that increases resiliency in the community. "
        "The persons to benefit from the Activities described herein must receive the prescribed service or benefit, and all eligibility requirements must be met to fulfill contractual obligations.")
    document.add_paragraph()

    para4 = document.add_paragraph(
        "The grant total is " + '${:,.2f}'.format((intake['CDBG-MIT Planned Amount'].values[0])) + ". Subrecipient will be required to maintain a detailed Budget breakdown in the official system of record "
        "(TIGR) of the GLO’s Community Development and Revitalization division. Subrecipient must ensure expenditures for individual projects do not exceed the amounts for detailed funding "
        "categories in the project budget of the approved Grant Application, as may be revised in writing upon mutual agreement of the Parties.  If it becomes necessary to redistribute Budget "
        "line-item amounts between Activities, Subrecipient must seek a Contract Amendment prior to performing any work.")
    document.add_paragraph()

    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #####################################################################################################################################################################################################

    Activity = readActivity(Intake)
    for i in range(len(Activity)):
        # add activity title for each activity
        para1 = document.add_paragraph(Activity['DRGR Activity Title'].values[i])
        para1.runs[0].font.underline = True
        para1.style.font.name = 'Times New Roman'
        para1.style.font.size = Pt(12)
        para1.style.font.bold = False
        document.add_paragraph()

        # add project description for each project under the same activity
        Project = readProject(Intake, Activity['Record ID#'].values[i])
        for j in range(len(Project)):
            para1 = document.add_paragraph(Project['Project: Project Title'].values[j] + ":")
            document.add_paragraph()
            para2 = para1.add_run(" Subrecipient shall " + Project['Project: Provide a detailed description of the scope of work proposed. For proposed work involving a length of road, ditch, channel, etc., report the scope of the project in linear feet (lf).'].values[j] +
                                     " Construction will take place at the following locations:")
            para2.style.font.name = 'Times New Roman'
            para2.style.font.size = Pt(12)
            para2.style.font.bold = False

            #####################################################################################################################

            # add project site table for each project under the same activity
            Project_Sites = readProjectSite(Intake, Activity['Record ID#'].values[i], Project['Record ID#'].values[j])
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells

            # Set a cell background (shading) color to RGB D9D9D9.
            shading0 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

            # set headers
            hdr_cells[0].text = 'Defining project location (on/along...)'
            run = hdr_cells[0].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading0)

            hdr_cells[1].text = 'Approximate path or location (from...to...) mid-point coordinates'
            run = hdr_cells[1].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading1)

            hdr_cells[2].text = 'Proposed HUD Performance Measures'
            run = hdr_cells[2].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[2]._tc.get_or_add_tcPr().append(shading2)

            # set cell values
            for k in range(len(Project_Sites)):
                row_cells = table.add_row().cells
                row_cells[0].text = Project_Sites['Site: Street Address'].values[k]
                row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[1].text = Project_Sites['Approximate path or  location (from.. to..)  mid-point coordinates'].values[k]
                row_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[2].text = Project_Sites['Proposed HUD Performance Measures'].values[k]
                row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

            document.add_paragraph()
        #####################################################################################################################

            # add Beneficiary table for each project under the same activity
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells

            # Set a cell background (shading) color to RGB D9D9D9.
            shading0 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading3 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            shading4 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

            # set headers
            hdr_cells[0].text = 'Total Beneficiaries'
            run = hdr_cells[0].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading0)

            hdr_cells[1].text = 'LMI Beneficiaries'
            run = hdr_cells[1].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading1)

            hdr_cells[2].text = 'LMI %'
            run = hdr_cells[2].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[2]._tc.get_or_add_tcPr().append(shading2)

            hdr_cells[3].text = 'Census Tract'
            run = hdr_cells[3].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[3]._tc.get_or_add_tcPr().append(shading3)

            hdr_cells[4].text = 'Block Group'
            run = hdr_cells[4].paragraphs[0].runs[0]
            run.font.bold = True
            hdr_cells[4].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[4]._tc.get_or_add_tcPr().append(shading4)

            # set cell values

            row_cells = table.add_row().cells
            row_cells[0].text = (str(int(Project['National: Provide Total Number of Beneficiaries.'].values[j])))
            row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[1].text = (str(int(Project['National: Provide number of LMI beneficiaries'].values[j])))
            row_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[2].text = (str('{:.2%}'.format(Project['National: Percentage of LMI Beneficiaries'].values[j])))
            row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[3].text = (str(Project['Census Tract'].values[j]))
            row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[4].text = [Project['Tract-Group (Combined)'].values[j][k].split("-")[1].replace(" ; ", ",") + "                                             " for k in range(len(Project['Tract-Group (Combined)'].values[j]))]
            row_cells[4].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

            exception = ""
            if Project["National: Is the applicant a HUD Exception Grantee?"].values[j] == "Yes":
                exception = word.capitalize() + intake['Applicant: Applicant'].values[0] + " is a HUD Exception Grantee. Their LMI threshold is " + (str('{:.2%}'.format(Project['National: Percentage of LMI Beneficiaries'].values[j])))

            table1 = document.add_table(rows=1, cols=1)
            table1.style = 'Table Grid'
            for k in range(len(Project)):
                table1.rows[0].cells[0].text = ("Beneficiaries were identified using " + Project['National: What method was used for Beneficiary Identification?'].values[k] + " and will meet the " +
                                           Project['National: Which HUD national objective does the project meet?'].values[k] + " national objective. " + exception)
            document.add_paragraph()
        # add footer
        document.add_paragraph()
        document.add_paragraph()
        para5 = document.add_paragraph("REMAINDER OF PAGE INTENTIONALLY LEFT BLANK")
        para5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Adding a page break
        document.add_page_break()

    #####################################################################################################################
    # add Budget header
    font_charstyle = font_styles.add_style('Budget Header', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(12)
    font_object.name = 'Times New Roman'
    font_object.bold = True

    para1 = document.add_paragraph()
    para1.alignment = WD_TABLE_ALIGNMENT.CENTER
    para1.add_run('BUDGET', style='Budget Header')

    # add Budget table for each activity
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    # get total row data
    tot1 = 0
    tot2 = 0
    tot3 = 0

    # get budget totals
    ga = 0
    env = 0
    eng = 0

    # Set a cell background (shading) color to RGB D9D9D9.
    shading0 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading3 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

    # set headers
    hdr_cells[0].text = 'DRGR Activity Type'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    hdr_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading0)

    hdr_cells[1].text = 'Grant Award'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    hdr_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1]._tc.get_or_add_tcPr().append(shading1)

    hdr_cells[2].text = 'Other Funds'
    run = hdr_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    hdr_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[2]._tc.get_or_add_tcPr().append(shading2)

    hdr_cells[3].text = 'Total'
    run = hdr_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    hdr_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[3]._tc.get_or_add_tcPr().append(shading3)

    for i in range(len(Activity)):
        row_cells = table.add_row().cells
        row_cells[0].text =str(Activity['DRGR Activity Title'].values[i])
        row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        row_cells[1].text = str("${:,.2f}".format(Activity['Grant Award'].values[i]))
        row_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        tot1 += Activity['Grant Award'].values[i]

        row_cells[2].text = str("${:,.2f}".format(Activity['Other Funds'].values[i]))
        row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        tot2 += Activity['Other Funds'].values[i]

        row_cells[3].text = str("${:,.2f}".format(Activity['Total'].values[i]))
        row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        tot3 += Activity['Total'].values[i]

        ga += Activity['Total Grant Admin'].values[i]
        env += Activity['Total Environmental'].values[i]
        eng += Activity['Total Engineering'].values[i]

    # add total row

    # Set a cell background (shading) color to RGB D9D9D9.
    shading1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading3 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    shading4 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

    row_cells = table.add_row().cells

    row_cells[0].text = "TOTAL"
    row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    run = row_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(11)
    row_cells[0]._tc.get_or_add_tcPr().append(shading1)

    row_cells[1].text = str("${:,.2f}".format(tot1))
    row_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    run = row_cells[1].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(11)
    row_cells[1]._tc.get_or_add_tcPr().append(shading2)

    row_cells[2].text = str("${:,.2f}".format(tot2))
    row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    run = row_cells[2].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(11)
    row_cells[2]._tc.get_or_add_tcPr().append(shading3)

    row_cells[3].text = str("${:,.2f}".format(tot3))
    row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    run = row_cells[3].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(11)
    row_cells[3]._tc.get_or_add_tcPr().append(shading4)

    # add Budget subtext
    font_charstyle = font_styles.add_style('Budget Subtext', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(10)
    font_object.name = 'Times New Roman'
    font_object.bold = True

    para2 = document.add_paragraph()
    para2.add_run("{}".format('\u00B9') + " " + (intake['Applicant: Applicant'].values[0]).upper() + " GENERAL FUND, TO BE USED FOR XXXX", style='Budget Subtext')

    document.add_page_break()

    #####################################################################################################################
    set_line_spacing(document)
    preventDocumentBreak(document)
    document.save('Draft 1.docx')

    master = Document("Draft 1.docx")

    composer = Composer(master)

    # filename_second_docx is the name of the second docx file
    doc2 = Document("Benchmarks Template.docx")

    # append the doc2 into the master using composer.append function
    composer.append(doc2)

    # Save the combined docx with a name
    os.chdir(finalDestination)
    composer.save("COG-MOD Attachment A - " + intake['Applicant: Applicant'].values[0] + " (" + intake['Ref ID'].values[0] + ").docx")

def download_file(recordID, local_filename):
    url = 'https://api.quickbase.com/v1/files/bssvkhm48/' + str(recordID) + '/74/0'
    headers = {
        'QB-Realm-Hostname': 'dmsrecovery.quickbase.com',
        'Authorization': 'QB-USER-TOKEN ' + 'b7n73q_g6xg_0_ttt8jjdvzpzzudurmt9vcryr5yz'
    }

    r = requests.get(url, stream=True, headers=headers)

    with open(local_filename, 'wb') as f:
        f.write(base64.decodebytes(r.content))

finalFolder = r'C:\Users\58741\OneDrive - ICF\Documents\Python Scripts\Attachment A\RFIs'
def downloadRFIs(Intake):
    os.chdir(finalFolder)
    merger = PdfMerger()
    pdfs = []

    RFIs = readQBData('bssvkhm48', '29', "Attachment: RFIs")
    RFIs.columns = RFIs.columns.str.lstrip()
    RFIs['Related Intake (Calc)'] = RFIs['Related Intake (Calc)'].astype(str).apply(lambda x: x.replace('.0', ''))
    filtered = RFIs[RFIs['Related Intake (Calc)'] == str(Intake)][['Related Intake (Calc)', 'Related Intake (Calc) - Applicant: Applicant', 'Related Intake (Calc) - Ref ID', 'Record ID#', 'EFP File Name']].reset_index()
    for row in range(len(filtered)):
        download_file(filtered.loc[row, 'Record ID#'], filtered.loc[row, 'EFP File Name'])
        pdfs.append(filtered.loc[row, 'EFP File Name'])

    for pdf in pdfs:
        merger.append(pdf)

    merger.write(filtered['Related Intake (Calc) - Applicant: Applicant'].values[0] + " (" + filtered['Related Intake (Calc) - Ref ID'].values[0] + ")" + ".pdf")
    merger.close()

    for pdf in pdfs:
        os.remove(pdf)

    return

createDoc(193)